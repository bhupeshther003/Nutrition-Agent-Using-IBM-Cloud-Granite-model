"""
Document processing utilities for extracting text from PDF and DOCX files
"""

import os
import PyPDF2
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Class for processing resume documents (PDF and DOCX)"""
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """
        Extract text from PDF file using multiple methods for better accuracy
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        text = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber didn't extract much, try PyPDF2
            if len(text.strip()) < 100:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX: {file_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path):
        """
        Extract text from document based on file extension
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def validate_file(filename):
        """
        Validate if file has allowed extension
        
        Args:
            filename (str): Name of the file
            
        Returns:
            bool: True if valid, False otherwise
        """
        allowed_extensions = {'pdf', 'docx'}
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in allowed_extensions
    
    @staticmethod
    def get_file_size(file_path):
        """
        Get file size in bytes
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            int: File size in bytes
        """
        return os.path.getsize(file_path)
    
    @staticmethod
    def clean_text(text):
        """
        Clean and normalize extracted text
        
        Args:
            text (str): Raw text
            
        Returns:
            str: Cleaned text
        """
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep important punctuation
        # This is a basic implementation - can be enhanced based on needs
        
        return text.strip()

# Made with Bob
